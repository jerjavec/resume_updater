from docx import Document
from docx2pdf import convert
from docx.shared import Pt
from datetime import datetime
from .logger_manager import log_update

def update_resume(docx_path, new_title, settings):
    """
    Updates the resume title in a Word document based on formatting or style detection,
    saves updated Word + PDF files, and logs the update.
    """
    doc = Document(docx_path)
    title_replaced = False

    # Change the 4th paragraph (index 3) to new_title with formatting
    if len(doc.paragraphs) < 4:
        raise ValueError("Document does not have enough paragraphs to update the title.")
    para = doc.paragraphs[3]
    para.clear()
    new_run = para.add_run(new_title)
    new_run.bold = True
    new_run.font.size = Pt(12)
    new_run.font.name = "Calibri"
    para.alignment = 1  # centered
    title_replaced = True

    import os
    today = datetime.today().strftime("%Y%m%d")
    safe_title = new_title.replace(" ", "-")
    name = settings["name"]
    input_dir = os.path.dirname(docx_path)
    new_docx_filename = os.path.join(input_dir, f"{today}-{name}_resume_{safe_title}.docx")
    pdf_filename = os.path.join(input_dir, settings["pdf_filename"])

    doc.save(new_docx_filename)
    try:
        convert(new_docx_filename, pdf_filename)
    except Exception as e:
        raise RuntimeError(f"PDF creation from {new_docx_filename} to {pdf_filename} failed: {e}")

    # Log this update
    log_update(docx_path, new_title, new_docx_filename, pdf_filename)

    return new_docx_filename, pdf_filename
