import pytest
import os
import sys
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'resume_updater')))
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'resume_updater/resume_updater')))
from resume_updater.logic import update_resume
from resume_updater.settings_manager import load_settings
from docx import Document
from docx.shared import Pt

def create_docx_with_title(path, title="Summary"):
    doc = Document()
    n = doc.add_paragraph("John Doe")  # name
    run = n.runs[0]
    run.bold = True
    run.font.size = Pt(12)  # simulate missing explicit size
    n.alignment = 1  # centered
    doc.add_paragraph("")          # blank
    doc.add_paragraph("")          # blank
    p = doc.add_paragraph(title)
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(12)  # simulate missing explicit size
    p.alignment = 1  # centered
    doc.save(path)

def test_update_resume_invalid_file(tmp_path):
    settings = load_settings()
    fake_file = tmp_path / "fake.docx"
    fake_file.write_text("Not a docx")
    with pytest.raises(Exception):
        update_resume(str(fake_file), "New Title", settings)

def test_update_resume_valid_file(tmp_path):
    settings = load_settings()
    file = tmp_path / "resume.docx"
    create_docx_with_title(file)
    new_docx, new_pdf = update_resume(str(file), "Manager", settings)
    assert os.path.exists(new_docx)
    assert os.path.exists(new_pdf)

def test_update_resume_changes_title(tmp_path):
    settings = load_settings()
    file = tmp_path / "resume.docx"
    create_docx_with_title(file, "Old Title")
    new_docx, _ = update_resume(str(file), "New Manager", settings)
    doc = Document(new_docx)
    assert "New Manager" in [p.text for p in doc.paragraphs]

def test_update_resume_no_title_found(tmp_path):
    settings = load_settings()
    file = tmp_path / "resume.docx"
    doc = Document()
    doc.add_paragraph("John Doe")
    doc.save(file)
    with pytest.raises(ValueError):
        update_resume(str(file), "Developer", settings)

def test_update_resume_with_style_override(tmp_path):
    settings = load_settings()
    settings["resume_title_style"] = "Heading 1"
    file = tmp_path / "resume.docx"
    doc = Document()
    p = doc.add_paragraph("Override Title")
    p.style = "Heading 1"
    p.alignment = 1
    doc.save(file)
    new_docx, _ = update_resume(str(file), "Director", settings)
    doc2 = Document(new_docx)
    assert "Director" in [p.text for p in doc2.paragraphs]
